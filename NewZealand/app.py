from __future__ import annotations

import re
import json
from html import escape
from pathlib import Path
from typing import Any

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

try:
    from docx import Document
except ImportError:  # pragma: no cover - Word support is optional
    Document = None


BASE_DIR = Path(__file__).resolve().parent
SUPPORTED_EXCEL = {".xlsx", ".xls", ".xlsm", ".csv"}
SUPPORTED_WORD = {".docx"}
URL_RE = re.compile(r"https?://[^\s<>\]\)\"']+", re.IGNORECASE)
SCHOOL_DOMAINS = {
    "University of Waikato": "waikato.ac.nz",
    "University of Otago": "otago.ac.nz",
    "University of Canterbury": "canterbury.ac.nz",
    "University of Auckland": "auckland.ac.nz",
    "Massey University": "massey.ac.nz",
    "Victoria University of Wellington": "wgtn.ac.nz",
}


st.set_page_config(
    page_title="文件浏览器",
    page_icon="📁",
    layout="wide",
)

st.markdown(
    """
    <style>
    [data-testid="stSidebar"] {
        background: #f7f8fa;
    }
    [data-testid="stSidebar"] .stButton > button {
        min-height: 118px;
        width: 100%;
        padding: 10px 8px;
        border-radius: 8px;
        border: 1px solid rgba(49, 51, 63, 0.16);
        background: rgba(255, 255, 255, 0.96);
        color: #31333f;
        line-height: 1.22;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        border-color: rgba(49, 51, 63, 0.42);
        background: #ffffff;
        color: #b42318;
    }
    [data-testid="stSidebar"] .stButton > button p {
        font-size: 0.9rem;
        overflow-wrap: anywhere;
        word-break: break-word;
        line-height: 1.2;
        margin: 0;
        white-space: normal;
    }
    [data-testid="stSidebar"] .file-marker {
        display: none;
    }
    [data-testid="stSidebar"] div[data-testid="stElementContainer"]:has(div[data-testid="stMarkdownContainer"] .excel-file-marker) + div[data-testid="stElementContainer"] .stButton > button::before {
        content: "🟩";
        display: block;
        font-size: 1.9rem;
        line-height: 1;
        margin-bottom: 12px;
    }
    [data-testid="stSidebar"] div[data-testid="stElementContainer"]:has(div[data-testid="stMarkdownContainer"] .word-file-marker) + div[data-testid="stElementContainer"] .stButton > button::before {
        content: "🟦";
        display: block;
        font-size: 1.9rem;
        line-height: 1;
        margin-bottom: 12px;
    }
    [data-testid="stSidebar"] div[data-testid="stElementContainer"]:has(div[data-testid="stMarkdownContainer"] .generic-file-marker) + div[data-testid="stElementContainer"] .stButton > button::before {
        content: "⬜";
        display: block;
        font-size: 1.9rem;
        line-height: 1;
        margin-bottom: 12px;
    }
    [data-testid="stSidebar"] div[data-testid="stElementContainer"]:has(div[data-testid="stMarkdownContainer"] .selected-file) + div[data-testid="stElementContainer"] .stButton > button {
        border-color: #d92d20;
        box-shadow: inset 0 0 0 1px #d92d20;
        background: #fff7f6;
        color: #d92d20;
    }
    .mentor-card-grid {
        display: grid;
        grid-template-columns: repeat(auto-fill, minmax(132px, 1fr));
        gap: 14px;
        margin-top: 12px;
    }
    .mentor-card {
        aspect-ratio: 1 / 1;
        border: 1px solid rgba(49, 51, 63, 0.16);
        border-radius: 8px;
        background: #ffffff;
        color: #31333f;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        gap: 9px;
        padding: 12px;
        text-align: center;
        text-decoration: none;
        box-shadow: 0 1px 2px rgba(16, 24, 40, 0.05);
        transition: border-color 120ms ease, transform 120ms ease, box-shadow 120ms ease;
    }
    .mentor-card:hover {
        border-color: #d92d20;
        color: #b42318;
        text-decoration: none;
        transform: translateY(-1px);
        box-shadow: 0 8px 18px rgba(16, 24, 40, 0.10);
    }
    .mentor-card.selected {
        border-color: #d92d20;
        box-shadow: inset 0 0 0 1px #d92d20;
        background: #fff7f6;
    }
    .mentor-logo-wrap {
        width: 58px;
        height: 58px;
        border-radius: 50%;
        display: grid;
        place-items: center;
        background: #f3f4f6;
        color: #475467;
        font-size: 0.8rem;
        font-weight: 700;
        overflow: hidden;
    }
    .mentor-logo-wrap::after {
        content: attr(data-fallback);
    }
    .mentor-logo {
        width: 44px;
        height: 44px;
        object-fit: contain;
        display: block;
    }
    .mentor-logo-wrap:has(.mentor-logo)::after {
        content: "";
    }
    .mentor-name {
        font-size: 0.9rem;
        font-weight: 700;
        line-height: 1.18;
    }
    .mentor-school {
        font-size: 0.72rem;
        line-height: 1.15;
        color: #667085;
    }
    .mentor-detail-panel {
        position: sticky;
        top: 16px;
        border: 1px solid rgba(49, 51, 63, 0.16);
        border-radius: 8px;
        background: #ffffff;
        padding: 16px;
        margin-top: 12px;
    }
    .mentor-detail-panel h3 {
        margin: 0 0 4px;
        font-size: 1.2rem;
        line-height: 1.25;
    }
    .mentor-detail-school {
        color: #667085;
        font-size: 0.86rem;
        margin-bottom: 12px;
    }
    .mentor-detail-row {
        margin-top: 10px;
        font-size: 0.88rem;
        line-height: 1.35;
    }
    .mentor-detail-row strong {
        display: block;
        color: #344054;
        margin-bottom: 2px;
    }
    .mentor-detail-row span {
        color: #475467;
    }
    .mentor-home-button {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        min-height: 38px;
        margin-top: 14px;
        padding: 0 16px;
        border-radius: 8px;
        background: #d92d20;
        color: #ffffff !important;
        text-decoration: none !important;
        font-weight: 700;
    }
    .mentor-home-button:hover {
        background: #b42318;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


def list_supported_files() -> list[Path]:
    files: list[Path] = []
    for path in BASE_DIR.iterdir():
        if not path.is_file() or path.name.startswith("~$"):
            continue
        if path.suffix.lower() in SUPPORTED_EXCEL | SUPPORTED_WORD:
            files.append(path)
    return sorted(files, key=lambda item: item.name.lower())


def file_type_icon(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in SUPPORTED_EXCEL:
        return "🟩 Excel"
    if suffix in SUPPORTED_WORD:
        return "🟦 Word"
    return "📁 File"


def large_file_icon(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in SUPPORTED_EXCEL:
        return "📄\n🟩 XLSX"
    if suffix in SUPPORTED_WORD:
        return "📄\n🟦 DOCX"
    return "📄\nFILE"


def compact_file_icon(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in SUPPORTED_EXCEL:
        return "🟩🟩🟩\nXLSX"
    if suffix in SUPPORTED_WORD:
        return "🟦🟦🟦\nDOCX"
    return "⬜⬜⬜\nFILE"


def file_square_icon(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in SUPPORTED_EXCEL:
        return "🟩"
    if suffix in SUPPORTED_WORD:
        return "🟦"
    return "⬜"


def file_marker_classes(path: Path, selected: bool) -> str:
    suffix = path.suffix.lower()
    classes = ["file-marker"]
    if suffix in SUPPORTED_EXCEL:
        classes.append("excel-file-marker")
    elif suffix in SUPPORTED_WORD:
        classes.append("word-file-marker")
    else:
        classes.append("generic-file-marker")
    if selected:
        classes.append("selected-file")
    return " ".join(classes)


def render_file_browser(files: list[Path]) -> Path:
    if "selected_file_name" not in st.session_state:
        st.session_state.selected_file_name = files[0].name

    file_names = {file.name for file in files}
    if st.session_state.selected_file_name not in file_names:
        st.session_state.selected_file_name = files[0].name

    with st.sidebar:
        st.header("文件")
        for row_start in range(0, len(files), 2):
            cols = st.columns(2)
            for col, file in zip(cols, files[row_start : row_start + 2]):
                selected = file.name == st.session_state.selected_file_name
                with col:
                    marker_classes = file_marker_classes(file, selected)
                    st.markdown(f'<span class="{marker_classes}"></span>', unsafe_allow_html=True)
                    if st.button(file.name, key=f"file_{file.name}", type="secondary"):
                        st.session_state.selected_file_name = file.name

    return next(file for file in files if file.name == st.session_state.selected_file_name)


def is_url(value: Any) -> bool:
    if pd.isna(value):
        return False
    text = str(value).strip()
    return bool(URL_RE.fullmatch(text))


def normalize_url(value: Any) -> str:
    text = "" if pd.isna(value) else str(value).strip()
    if not text:
        return ""
    if text.startswith(("http://", "https://")):
        return text
    if "." in text and " " not in text:
        return f"https://{text}"
    return text


def school_logo_url(school: str) -> str:
    domain = SCHOOL_DOMAINS.get(school.strip())
    if not domain:
        return ""
    return f"https://www.google.com/s2/favicons?sz=128&domain={domain}"


def school_initials(school: str) -> str:
    words = [word for word in re.split(r"\W+", school) if word and word.lower() not in {"of", "the"}]
    return "".join(word[0].upper() for word in words[:2]) or "NZ"


def school_css_class(school: str) -> str:
    school = school.strip()
    mapping = {
        "University of Waikato": "school-waikato",
        "University of Otago": "school-otago",
        "University of Canterbury": "school-canterbury",
        "University of Auckland": "school-auckland",
        "Massey University": "school-massey",
        "Victoria University of Wellington": "school-wellington",
    }
    return mapping.get(school, "school-generic")


def find_column(df: pd.DataFrame, candidates: list[str]) -> str | None:
    lowered = {str(col).lower(): col for col in df.columns}
    for candidate in candidates:
        candidate_lower = candidate.lower()
        for lower_name, original in lowered.items():
            if candidate_lower in lower_name:
                return original
    return None


def value_text(row: pd.Series, col: str | None) -> str:
    if not col:
        return ""
    value = row.get(col, "")
    if pd.isna(value):
        return ""
    return str(value).strip()


def mentor_explorer_html(mentors: list[dict[str, Any]]) -> str:
    payload = json.dumps(mentors, ensure_ascii=False)
    return f"""
<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{
    margin: 0;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    color: #31333f;
}}
.mentor-layout {{
    display: grid;
    grid-template-columns: minmax(0, 0.62fr) minmax(280px, 0.38fr);
    gap: 24px;
    align-items: start;
}}
.school-tabs {{
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
    margin: 0 0 14px;
}}
.school-tab {{
    border: 1px solid rgba(49, 51, 63, 0.16);
    border-radius: 8px;
    background: #fff;
    color: #344054;
    cursor: pointer;
    font-size: 0.82rem;
    font-weight: 700;
    padding: 8px 10px;
}}
.school-tab:hover,
.school-tab.active {{
    border-color: #d92d20;
    color: #b42318;
    background: #fff7f6;
}}
.school-summary {{
    margin: 0 0 12px;
    color: #667085;
    font-size: 0.86rem;
}}
.mentor-card-grid {{
    display: grid;
    grid-template-columns: repeat(auto-fill, minmax(132px, 1fr));
    gap: 14px;
}}
.mentor-card {{
    aspect-ratio: 1 / 1;
    border: 1px solid rgba(49, 51, 63, 0.16);
    border-radius: 8px;
    background: #fff;
    color: #31333f;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    gap: 9px;
    padding: 12px;
    text-align: center;
    text-decoration: none;
    box-shadow: 0 1px 2px rgba(16, 24, 40, 0.05);
    cursor: pointer;
}}
.mentor-card:hover {{
    border-color: #d92d20;
    color: #b42318;
    transform: translateY(-1px);
    box-shadow: 0 8px 18px rgba(16, 24, 40, 0.10);
}}
.mentor-card.selected {{
    border-color: #d92d20;
    box-shadow: inset 0 0 0 1px #d92d20;
    background: #fff7f6;
}}
.mentor-logo-wrap {{
    width: 58px;
    height: 58px;
    border-radius: 50%;
    display: grid;
    place-items: center;
    background: #f3f4f6;
    color: #475467;
    font-size: 0.8rem;
    font-weight: 700;
    overflow: hidden;
}}
.mentor-logo {{
    width: 44px;
    height: 44px;
    object-fit: contain;
    display: block;
}}
.mentor-name {{
    font-size: 0.9rem;
    font-weight: 700;
    line-height: 1.18;
}}
.mentor-school {{
    font-size: 0.72rem;
    line-height: 1.15;
    color: #667085;
}}
.mentor-detail-panel {{
    position: sticky;
    top: 0;
    border: 1px solid rgba(49, 51, 63, 0.16);
    border-radius: 8px;
    background: #fff;
    padding: 16px;
}}
.mentor-detail-panel h3 {{
    margin: 0 0 4px;
    font-size: 1.2rem;
    line-height: 1.25;
}}
.mentor-detail-school {{
    color: #667085;
    font-size: 0.86rem;
    margin-bottom: 12px;
}}
.mentor-detail-row {{
    margin-top: 10px;
    font-size: 0.88rem;
    line-height: 1.35;
}}
.mentor-detail-row strong {{
    display: block;
    color: #344054;
    margin-bottom: 2px;
}}
.mentor-detail-row span {{
    color: #475467;
}}
.mentor-home-button {{
    display: inline-flex;
    align-items: center;
    justify-content: center;
    min-height: 38px;
    margin-top: 14px;
    padding: 0 16px;
    border-radius: 8px;
    background: #d92d20;
    color: #fff;
    text-decoration: none;
    font-weight: 700;
}}
.mentor-home-button:hover {{
    background: #b42318;
}}
@media (max-width: 760px) {{
    .mentor-layout {{
        grid-template-columns: 1fr;
    }}
    .mentor-detail-panel {{
        position: static;
    }}
}}
</style>
</head>
<body>
<nav id="schoolTabs" class="school-tabs" aria-label="学校分页"></nav>
<p id="schoolSummary" class="school-summary"></p>
<div class="mentor-layout">
    <div id="cards" class="mentor-card-grid"></div>
    <aside id="detail" class="mentor-detail-panel"></aside>
</div>
<script>
const mentors = {payload};
const schools = [...new Set(mentors.map(mentor => mentor.school || "未标注学校"))];
const schoolTabs = document.getElementById("schoolTabs");
const schoolSummary = document.getElementById("schoolSummary");
const cards = document.getElementById("cards");
const detail = document.getElementById("detail");
let selectedSchool = schools[0] || "";
let selectedIndex = 0;

function escapeHtml(value) {{
    return String(value ?? "").replace(/[&<>"']/g, ch => ({{
        "&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;", "'": "&#39;"
    }}[ch]));
}}

function renderCards() {{
    const visibleMentors = mentors
        .map((mentor, index) => ({{...mentor, index}}))
        .filter(mentor => (mentor.school || "未标注学校") === selectedSchool);
    schoolSummary.textContent = `${{selectedSchool}} · ${{visibleMentors.length}} 位导师 · 按学校内投递优先度排序`;
    cards.innerHTML = visibleMentors.map((mentor) => `
        <button class="mentor-card ${{mentor.index === selectedIndex ? "selected" : ""}}" type="button" data-index="${{mentor.index}}">
            <span class="mentor-logo-wrap">
                ${{mentor.logo ? `<img class="mentor-logo" src="${{escapeHtml(mentor.logo)}}" alt="${{escapeHtml(mentor.school)}} logo" onerror="this.replaceWith(document.createTextNode('${{escapeHtml(mentor.initials)}}'));">` : escapeHtml(mentor.initials)}}
            </span>
            <span class="mentor-name">${{escapeHtml(mentor.name)}}</span>
            <span class="mentor-school">${{escapeHtml(mentor.school)}}</span>
        </button>
    `).join("");
}}

function renderDetail() {{
    const mentor = mentors[selectedIndex];
    if (!mentor) {{
        detail.innerHTML = "";
        return;
    }}
    const rows = mentor.facts
        .filter(item => item.value)
        .map(item => `<div class="mentor-detail-row"><strong>${{escapeHtml(item.label)}}</strong><span>${{escapeHtml(item.value)}}</span></div>`)
        .join("");
    detail.innerHTML = `
        <h3>${{escapeHtml(mentor.name)}}</h3>
        <div class="mentor-detail-school">${{escapeHtml(mentor.school)}}</div>
        ${{rows}}
        <a class="mentor-home-button" href="${{escapeHtml(mentor.url)}}" target="_blank" rel="noopener noreferrer">主页</a>
    `;
}}

function renderSchoolTabs() {{
    schoolTabs.innerHTML = schools.map(school => `
        <button class="school-tab ${{school === selectedSchool ? "active" : ""}}" type="button" data-school="${{escapeHtml(school)}}">
            ${{escapeHtml(school)}} <span>${{mentors.filter(mentor => (mentor.school || "未标注学校") === school).length}}</span>
        </button>
    `).join("");
}}

cards.addEventListener("click", event => {{
    const card = event.target.closest(".mentor-card");
    if (!card) return;
    selectedIndex = Number(card.dataset.index);
    renderCards();
    renderDetail();
}});

schoolTabs.addEventListener("click", event => {{
    const tab = event.target.closest(".school-tab");
    if (!tab) return;
    selectedSchool = tab.dataset.school;
    const firstMentorIndex = mentors.findIndex(mentor => (mentor.school || "未标注学校") === selectedSchool);
    selectedIndex = Math.max(firstMentorIndex, 0);
    renderSchoolTabs();
    renderCards();
    renderDetail();
}});

mentors.sort((left, right) => {{
    const schoolOrder = schools.indexOf(left.school || "未标注学校") - schools.indexOf(right.school || "未标注学校");
    if (schoolOrder !== 0) return schoolOrder;
    return Number(left.schoolRank || 9999) - Number(right.schoolRank || 9999);
}});
selectedSchool = schools[0] || "";
selectedIndex = mentors.findIndex(mentor => (mentor.school || "未标注学校") === selectedSchool);
if (selectedIndex < 0) selectedIndex = 0;
renderSchoolTabs();
renderCards();
renderDetail();
</script>
</body>
</html>
"""


def clean_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    df = df.dropna(how="all").dropna(axis=1, how="all")
    df = df.loc[:, ~df.columns.astype(str).str.match(r"^Unnamed")]
    df = df.reset_index(drop=True)
    df.columns = [str(col).strip() for col in df.columns]
    return df


def cell_text(value: Any) -> str:
    if pd.isna(value):
        return ""
    return str(value).strip()


def deduplicate_columns(columns: list[str]) -> list[str]:
    seen: dict[str, int] = {}
    deduped: list[str] = []
    for index, column in enumerate(columns, start=1):
        name = column or f"Column {index}"
        count = seen.get(name, 0)
        seen[name] = count + 1
        deduped.append(name if count == 0 else f"{name}_{count + 1}")
    return deduped


def header_row_score(raw: pd.DataFrame, row_index: int) -> float:
    row = raw.iloc[row_index]
    labels = [cell_text(value) for value in row]
    non_empty = [label for label in labels if label]
    if len(non_empty) < 2:
        return -1

    following = raw.iloc[row_index + 1 : row_index + 6]
    following_width = 0
    if not following.empty:
        following_width = int(following.notna().sum(axis=1).max())
    if following_width < 2:
        return -1

    duplicate_penalty = len(non_empty) - len(set(non_empty))
    long_text_penalty = sum(len(label) > 80 for label in non_empty)
    numeric_penalty = sum(label.replace(".", "", 1).isdigit() for label in non_empty)
    return (len(non_empty) * 3) + following_width - duplicate_penalty - long_text_penalty - numeric_penalty


def detect_header_row(raw: pd.DataFrame) -> int:
    scan_rows = min(12, len(raw))
    scored_rows = [(header_row_score(raw, index), index) for index in range(scan_rows)]
    valid_rows = [item for item in scored_rows if item[0] >= 0]
    if not valid_rows:
        return 0
    return max(valid_rows, key=lambda item: (item[0], -item[1]))[1]


def dataframe_from_raw(raw: pd.DataFrame) -> pd.DataFrame:
    raw = raw.dropna(how="all").dropna(axis=1, how="all")
    if raw.empty:
        return raw

    header_index = detect_header_row(raw)
    columns = deduplicate_columns([cell_text(value) for value in raw.iloc[header_index]])
    data = raw.iloc[header_index + 1 :].copy()
    data.columns = columns
    return clean_dataframe(data)


@st.cache_data(show_spinner=False)
def read_excel_file(path_text: str, sheet_name: str | int | None = None) -> pd.DataFrame:
    path = Path(path_text)
    if path.suffix.lower() == ".csv":
        return dataframe_from_raw(pd.read_csv(path, header=None))
    return dataframe_from_raw(pd.read_excel(path, sheet_name=sheet_name, header=None))


@st.cache_data(show_spinner=False)
def get_excel_sheets(path_text: str) -> list[str]:
    path = Path(path_text)
    if path.suffix.lower() == ".csv":
        return ["CSV"]
    return pd.ExcelFile(path).sheet_names


def likely_url_columns(df: pd.DataFrame) -> list[str]:
    columns: list[str] = []
    for col in df.columns:
        name = str(col).lower()
        sample = df[col].dropna().head(30)
        url_count = sum(is_url(value) for value in sample)
        if "url" in name or "link" in name or "主页" in str(col) or url_count > 0:
            columns.append(col)
    return columns


def likely_person_column(df: pd.DataFrame) -> str:
    for col in df.columns:
        name = str(col).lower()
        if any(token in name for token in ["mentor", "name", "supervisor", "导师", "姓名"]):
            return col
    if len(df.columns) >= 3:
        return df.columns[2]
    return df.columns[0]


def show_dataframe(df: pd.DataFrame) -> None:
    url_cols = likely_url_columns(df)
    column_config = {
        col: st.column_config.LinkColumn(col, display_text="打开链接")
        for col in url_cols
    }
    st.dataframe(
        df,
        use_container_width=True,
        hide_index=True,
        column_config=column_config,
    )


def show_mentor_links(df: pd.DataFrame, file_name: str) -> None:
    column_text = " ".join(str(col).lower() for col in df.columns)
    looks_like_mentor_file = "mentor" in file_name.lower() or "导师" in column_text
    if not looks_like_mentor_file:
        return

    url_cols = likely_url_columns(df)
    if not url_cols or df.empty:
        return

    url_col = url_cols[0]
    name_col = likely_person_column(df)
    school_col = next(
        (col for col in df.columns if "school" in str(col).lower() or "学校" in str(col)),
        None,
    )
    position_col = find_column(df, ["职位", "position"])
    research_col = find_column(df, ["研究方向", "research"])
    match_col = find_column(df, ["Bioinformatics", "匹配"])
    phd_col = find_column(df, ["PhD"])
    ra_col = find_column(df, ["RA"])
    funding_col = find_column(df, ["funding"])
    note_col = find_column(df, ["联系", "备注", "note"])
    school_rank_col = find_column(df, ["学校内投递排名", "校内", "投递排名"])
    school_priority_col = find_column(df, ["学校投递优先度", "优先度分"])
    confirm_col = find_column(df, ["Bioinfo确认", "确认"])

    rows = df[df[url_col].apply(is_url)].copy()
    if rows.empty:
        return

    st.divider()
    st.subheader("导师主页快捷入口")
    st.caption(f"点击导师名片会在本页显示重要信息；需要离开时再点信息卡里的“主页”。按当前 Excel 表格推荐顺序显示全部 {len(rows)} 位导师。")

    mentors: list[dict[str, Any]] = []
    for _, row in rows.iterrows():
        name = str(row.get(name_col, "导师")).strip()
        school = str(row.get(school_col, "")).strip() if school_col else ""
        mentors.append(
            {
                "name": name,
                "school": school,
                "schoolRank": value_text(row, school_rank_col),
                "url": normalize_url(row[url_col]),
                "logo": school_logo_url(school),
                "initials": school_initials(school),
                "facts": [
                    {"label": "学校内投递排名", "value": value_text(row, school_rank_col)},
                    {"label": "学校投递优先度分", "value": value_text(row, school_priority_col)},
                    {"label": "Bioinfo 确认状态", "value": value_text(row, confirm_col)},
                    {"label": "职位/职称", "value": value_text(row, position_col)},
                    {"label": "研究方向", "value": value_text(row, research_col)},
                    {"label": "Bioinformatics 匹配", "value": value_text(row, match_col)},
                    {"label": "PhD 机会", "value": value_text(row, phd_col)},
                    {"label": "RA 机会", "value": value_text(row, ra_col)},
                    {"label": "Funding / 经费", "value": value_text(row, funding_col)},
                    {"label": "联系建议", "value": value_text(row, note_col)},
                ],
            }
        )

    components.html(mentor_explorer_html(mentors), height=900, scrolling=True)


def show_excel(path: Path) -> None:
    sheets = get_excel_sheets(str(path))
    selected_sheet = sheets[0]
    if len(sheets) > 1:
        selected_sheet = st.radio(
            "工作表",
            sheets,
            index=0,
            horizontal=True,
        )
    else:
        st.caption(f"工作表：{selected_sheet}")

    df = read_excel_file(str(path), None if selected_sheet == "CSV" else selected_sheet)
    if df.empty:
        st.warning("这个工作表没有可展示的数据。")
        return

    st.caption(f"{path.name} / {selected_sheet} · {len(df):,} 行 · {len(df.columns):,} 列")
    show_dataframe(df)
    show_mentor_links(df, path.name)


@st.cache_data(show_spinner=False)
def read_word(path_text: str) -> list[dict[str, Any]]:
    if Document is None:
        return [{"type": "error", "text": "未安装 python-docx，无法读取 Word 文件。"}]

    doc = Document(path_text)
    blocks: list[dict[str, Any]] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append({"type": "paragraph", "text": text})
    for table in doc.tables:
        table_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        blocks.append({"type": "table", "rows": table_rows})
    return blocks


def show_word(path: Path) -> None:
    blocks = read_word(str(path))
    for block in blocks:
        if block["type"] == "error":
            st.error(block["text"])
        elif block["type"] == "paragraph":
            st.write(block["text"])
        elif block["type"] == "table":
            rows = block["rows"]
            if rows:
                header, data = rows[0], rows[1:]
                st.dataframe(pd.DataFrame(data, columns=header), use_container_width=True)


def main() -> None:
    st.title("文件浏览器")
    st.caption("左侧选择文件，右侧查看工作表和内容。")

    files = list_supported_files()
    if not files:
        st.warning("当前文件夹下没有可展示的 Excel 或 Word 文件。")
        return

    selected = render_file_browser(files)
    st.header(selected.name)

    try:
        if selected.suffix.lower() in SUPPORTED_EXCEL:
            show_excel(selected)
        elif selected.suffix.lower() in SUPPORTED_WORD:
            show_word(selected)
        else:
            st.info("暂不支持这个文件类型。")
    except Exception as exc:
        st.exception(exc)


if __name__ == "__main__":
    main()

